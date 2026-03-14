#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter
将Markdown文本转换为Word文档的全面转换器
"""

import re
import os
import html
from typing import List, Tuple, Optional, Dict, Any
from dataclasses import dataclass
from enum import Enum, auto

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class ElementType(Enum):
    HEADING = auto()
    PARAGRAPH = auto()
    BOLD = auto()
    ITALIC = auto()
    STRIKETHROUGH = auto()
    CODE = auto()
    CODE_BLOCK = auto()
    LINK = auto()
    IMAGE = auto()
    LIST_ORDERED = auto()
    LIST_UNORDERED = auto()
    LIST_TASK = auto()
    BLOCKQUOTE = auto()
    TABLE = auto()
    HORIZONTAL_RULE = auto()
    FOOTNOTE = auto()
    MATH_INLINE = auto()
    MATH_BLOCK = auto()


@dataclass
class TextRun:
    text: str
    bold: bool = False
    italic: bool = False
    strikethrough: bool = False
    code: bool = False
    link: Optional[str] = None
    font_size: Optional[int] = None
    font_color: Optional[Tuple[int, int, int]] = None


@dataclass
class MarkdownElement:
    element_type: ElementType
    content: Any
    level: int = 0
    language: str = ""
    ordered: bool = False
    checked: bool = False
    rows: int = 0
    cols: int = 0


class MarkdownParser:
    """Markdown解析器"""

    def __init__(self):
        self.footnotes: Dict[str, str] = {}
        self.footnote_refs: List[str] = []

    def parse(self, markdown_text: str) -> List[MarkdownElement]:
        elements = []
        lines = markdown_text.split("\n")
        i = 0

        self._extract_footnotes(lines)

        while i < len(lines):
            line = lines[i]

            if self._is_empty_line(line):
                i += 1
                continue

            if self._is_code_block_start(line):
                element, i = self._parse_code_block(lines, i)
                elements.append(element)
                continue

            if self._is_table_line(line) and i + 1 < len(lines) and "|" in lines[i + 1]:
                element, i = self._parse_table(lines, i)
                elements.append(element)
                continue

            if self._is_horizontal_rule(line):
                elements.append(MarkdownElement(ElementType.HORIZONTAL_RULE, None))
                i += 1
                continue

            if self._is_heading(line):
                elements.append(self._parse_heading(line))
                i += 1
                continue

            if self._is_blockquote(line):
                element, i = self._parse_blockquote(lines, i)
                elements.append(element)
                continue

            if self._is_list_item(line):
                element, i = self._parse_list(lines, i)
                elements.append(element)
                continue

            if self._is_math_block_start(line):
                element, i = self._parse_math_block(lines, i)
                elements.append(element)
                continue

            element, i = self._parse_paragraph(lines, i)
            elements.append(element)

        return elements

    def _extract_footnotes(self, lines: List[str]):
        footnote_pattern = r"^\[\^([^\]]+)\]:\s*(.+)$"
        i = 0
        while i < len(lines):
            match = re.match(footnote_pattern, lines[i])
            if match:
                footnote_id = match.group(1)
                footnote_text = match.group(2)
                self.footnotes[footnote_id] = footnote_text
                lines.pop(i)
            else:
                i += 1

    def _is_empty_line(self, line: str) -> bool:
        return line.strip() == ""

    def _is_heading(self, line: str) -> bool:
        return line.strip().startswith("#")

    def _parse_heading(self, line: str) -> MarkdownElement:
        line = line.strip()
        level = 0
        while level < len(line) and line[level] == "#":
            level += 1
        content = line[level:].strip()
        runs = self._parse_inline(content)
        return MarkdownElement(ElementType.HEADING, runs, level=level)

    def _is_code_block_start(self, line: str) -> bool:
        return line.strip().startswith("```")

    def _parse_code_block(
        self, lines: List[str], start: int
    ) -> Tuple[MarkdownElement, int]:
        first_line = lines[start].strip()
        language = first_line[3:].strip() if len(first_line) > 3 else ""

        code_lines = []
        i = start + 1
        while i < len(lines):
            if lines[i].strip().startswith("```"):
                break
            code_lines.append(lines[i])
            i += 1

        code_content = "\n".join(code_lines)
        return (
            MarkdownElement(ElementType.CODE_BLOCK, code_content, language=language),
            i + 1,
        )

    def _is_table_line(self, line: str) -> bool:
        return line.strip().startswith("|") and line.strip().endswith("|")

    def _parse_table(self, lines: List[str], start: int) -> Tuple[MarkdownElement, int]:
        table_lines = []
        i = start
        while i < len(lines) and self._is_table_line(lines[i]):
            table_lines.append(lines[i])
            i += 1

        if len(table_lines) < 2:
            return (
                MarkdownElement(
                    ElementType.PARAGRAPH, self._parse_inline(lines[start])
                ),
                start + 1,
            )

        header_cells = self._parse_table_row(table_lines[0])
        separator_line = table_lines[1]
        alignments = self._parse_table_alignment(separator_line)

        data_rows = []
        for line in table_lines[2:]:
            data_rows.append(self._parse_table_row(line))

        table_data = {
            "header": header_cells,
            "rows": data_rows,
            "alignments": alignments,
        }

        return MarkdownElement(ElementType.TABLE, table_data), i

    def _parse_table_row(self, line: str) -> List[str]:
        cells = line.strip().split("|")
        return [cell.strip() for cell in cells[1:-1]]

    def _parse_table_alignment(self, line: str) -> List[str]:
        cells = line.strip().split("|")
        alignments = []
        for cell in cells[1:-1]:
            cell = cell.strip()
            if cell.startswith(":") and cell.endswith(":"):
                alignments.append("center")
            elif cell.endswith(":"):
                alignments.append("right")
            else:
                alignments.append("left")
        return alignments

    def _is_horizontal_rule(self, line: str) -> bool:
        stripped = line.strip()
        patterns = [r"^-{3,}$", r"^\*{3,}$", r"^_{3,}$"]
        return any(re.match(p, stripped) for p in patterns)

    def _is_blockquote(self, line: str) -> bool:
        return line.strip().startswith(">")

    def _parse_blockquote(
        self, lines: List[str], start: int
    ) -> Tuple[MarkdownElement, int]:
        quote_lines = []
        i = start
        while i < len(lines) and self._is_blockquote(lines[i]):
            quote_lines.append(lines[i].strip()[1:].strip())
            i += 1

        return MarkdownElement(ElementType.BLOCKQUOTE, quote_lines), i

    def _is_list_item(self, line: str) -> bool:
        stripped = line.strip()
        ordered_pattern = r"^\d+\.\s"
        unordered_pattern = r"^[-*+]\s"
        task_pattern = r"^[-*+]\s\[[ xX]\]\s"
        return bool(
            re.match(ordered_pattern, stripped)
            or re.match(unordered_pattern, stripped)
            or re.match(task_pattern, stripped)
        )

    def _get_list_indent(self, line: str) -> int:
        indent = 0
        for char in line:
            if char == " ":
                indent += 1
            elif char == "\t":
                indent += 4
            else:
                break
        return indent

    def _parse_list(self, lines: List[str], start: int) -> Tuple[MarkdownElement, int]:
        items = []
        i = start
        ordered = False
        is_task_list = False
        base_indent = self._get_list_indent(lines[start])

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            current_indent = self._get_list_indent(line)

            if self._is_empty_line(stripped):
                if i + 1 < len(lines) and self._is_list_item(lines[i + 1]):
                    i += 1
                    continue
                else:
                    break

            if not self._is_list_item(stripped):
                break

            ordered_match = re.match(r"^(\d+)\.\s(.+)$", stripped)
            task_match = re.match(r"^[-*+]\s\[([ xX])\]\s(.+)$", stripped)
            unordered_match = re.match(r"^[-*+]\s(.+)$", stripped)

            if ordered_match:
                ordered = True
                items.append(
                    {
                        "content": self._parse_inline(ordered_match.group(2)),
                        "checked": None,
                        "indent": max(0, (current_indent - base_indent) // 2),
                    }
                )
            elif task_match:
                is_task_list = True
                checked = task_match.group(1).lower() == "x"
                items.append(
                    {
                        "content": self._parse_inline(task_match.group(2)),
                        "checked": checked,
                        "indent": max(0, (current_indent - base_indent) // 2),
                    }
                )
            elif unordered_match:
                items.append(
                    {
                        "content": self._parse_inline(unordered_match.group(1)),
                        "checked": None,
                        "indent": max(0, (current_indent - base_indent) // 2),
                    }
                )
            i += 1

        element_type = (
            ElementType.LIST_TASK
            if is_task_list
            else (ElementType.LIST_ORDERED if ordered else ElementType.LIST_UNORDERED)
        )
        return MarkdownElement(element_type, items, ordered=ordered), i

    def _is_math_block_start(self, line: str) -> bool:
        return line.strip().startswith("$$")

    def _parse_math_block(
        self, lines: List[str], start: int
    ) -> Tuple[MarkdownElement, int]:
        math_lines = []
        i = start + 1
        while i < len(lines):
            if lines[i].strip().startswith("$$"):
                break
            math_lines.append(lines[i])
            i += 1

        math_content = "\n".join(math_lines)
        return MarkdownElement(ElementType.MATH_BLOCK, math_content), i + 1

    def _parse_paragraph(
        self, lines: List[str], start: int
    ) -> Tuple[MarkdownElement, int]:
        paragraph_lines = []
        i = start
        while i < len(lines):
            line = lines[i]
            if (
                self._is_empty_line(line)
                or self._is_heading(line)
                or self._is_code_block_start(line)
                or self._is_list_item(line)
                or self._is_blockquote(line)
                or self._is_horizontal_rule(line)
                or self._is_table_line(line)
            ):
                break
            paragraph_lines.append(line)
            i += 1

        content = " ".join(paragraph_lines)
        runs = self._parse_inline(content)
        return MarkdownElement(ElementType.PARAGRAPH, runs), i

    def _parse_inline(self, text: str) -> List[TextRun]:
        runs = []
        i = 0
        current_text = ""
        current_bold = False
        current_italic = False
        current_strikethrough = False
        current_code = False
        current_link = None

        while i < len(text):
            if text[i : i + 2] == "**" or text[i : i + 2] == "__":
                if current_text:
                    runs.append(
                        TextRun(
                            text=current_text,
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=current_link,
                        )
                    )
                    current_text = ""
                current_bold = not current_bold
                i += 2
                continue

            if text[i] == "*" or text[i] == "_":
                if i + 1 < len(text) and (text[i + 1] == "*" or text[i + 1] == "_"):
                    i += 1
                    continue
                if current_text:
                    runs.append(
                        TextRun(
                            text=current_text,
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=current_link,
                        )
                    )
                    current_text = ""
                current_italic = not current_italic
                i += 1
                continue

            if text[i : i + 2] == "~~":
                if current_text:
                    runs.append(
                        TextRun(
                            text=current_text,
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=current_link,
                        )
                    )
                    current_text = ""
                current_strikethrough = not current_strikethrough
                i += 2
                continue

            if text[i] == "`" and not current_code:
                if current_text:
                    runs.append(
                        TextRun(
                            text=current_text,
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=current_link,
                        )
                    )
                    current_text = ""
                current_code = True
                i += 1
                continue
            elif text[i] == "`" and current_code:
                current_code = False
                i += 1
                continue

            if text[i] == "[":
                link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", text[i:])
                if link_match:
                    if current_text:
                        runs.append(
                            TextRun(
                                text=current_text,
                                bold=current_bold,
                                italic=current_italic,
                                strikethrough=current_strikethrough,
                                code=current_code,
                                link=current_link,
                            )
                        )
                        current_text = ""

                    link_text = link_match.group(1)
                    link_url = link_match.group(2)
                    runs.append(
                        TextRun(
                            text=link_text,
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=link_url,
                        )
                    )
                    i += len(link_match.group(0))
                    continue

            if text[i : i + 2] == "![":
                img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", text[i:])
                if img_match:
                    if current_text:
                        runs.append(
                            TextRun(
                                text=current_text,
                                bold=current_bold,
                                italic=current_italic,
                                strikethrough=current_strikethrough,
                                code=current_code,
                                link=current_link,
                            )
                        )
                        current_text = ""

                    img_alt = img_match.group(1)
                    img_url = img_match.group(2)
                    runs.append(
                        TextRun(
                            text=f"[IMAGE:{img_url}:{img_alt}]",
                            bold=current_bold,
                            italic=current_italic,
                            strikethrough=current_strikethrough,
                            code=current_code,
                            link=None,
                        )
                    )
                    i += len(img_match.group(0))
                    continue

            footnote_match = re.match(r"\[\^([^\]]+)\]", text[i:])
            if footnote_match:
                footnote_id = footnote_match.group(1)
                if footnote_id in self.footnotes:
                    if current_text:
                        runs.append(
                            TextRun(
                                text=current_text,
                                bold=current_bold,
                                italic=current_italic,
                                strikethrough=current_strikethrough,
                                code=current_code,
                                link=current_link,
                            )
                        )
                        current_text = ""

                    runs.append(
                        TextRun(
                            text=f"[^{footnote_id}]",
                            bold=True,
                            italic=False,
                            strikethrough=False,
                            code=False,
                            link=None,
                        )
                    )
                    i += len(footnote_match.group(0))
                    continue

            if text[i : i + 2] == "$$":
                math_match = re.match(r"\$\$(.+?)\$\$", text[i:])
                if math_match:
                    if current_text:
                        runs.append(
                            TextRun(
                                text=current_text,
                                bold=current_bold,
                                italic=current_italic,
                                strikethrough=current_strikethrough,
                                code=current_code,
                                link=current_link,
                            )
                        )
                        current_text = ""

                    math_content = math_match.group(1)
                    runs.append(
                        TextRun(
                            text=math_content,
                            bold=False,
                            italic=True,
                            strikethrough=False,
                            code=True,
                            link=None,
                        )
                    )
                    i += len(math_match.group(0))
                    continue

            if text[i] == "$":
                math_match = re.match(r"\$([^$]+)\$", text[i:])
                if math_match:
                    if current_text:
                        runs.append(
                            TextRun(
                                text=current_text,
                                bold=current_bold,
                                italic=current_italic,
                                strikethrough=current_strikethrough,
                                code=current_code,
                                link=current_link,
                            )
                        )
                        current_text = ""

                    math_content = math_match.group(1)
                    runs.append(
                        TextRun(
                            text=math_content,
                            bold=False,
                            italic=True,
                            strikethrough=False,
                            code=True,
                            link=None,
                        )
                    )
                    i += len(math_match.group(0))
                    continue

            if text[i] == "\\" and i + 1 < len(text):
                current_text += text[i + 1]
                i += 2
                continue

            current_text += text[i]
            i += 1

        if current_text:
            runs.append(
                TextRun(
                    text=current_text,
                    bold=current_bold,
                    italic=current_italic,
                    strikethrough=current_strikethrough,
                    code=current_code,
                    link=current_link,
                )
            )

        return self._merge_runs(runs)

    def _merge_runs(self, runs: List[TextRun]) -> List[TextRun]:
        if not runs:
            return runs

        merged = [runs[0]]
        for run in runs[1:]:
            last = merged[-1]
            if (
                last.bold == run.bold
                and last.italic == run.italic
                and last.strikethrough == run.strikethrough
                and last.code == run.code
                and last.link == run.link
            ):
                last.text += run.text
            else:
                merged.append(run)

        return merged


class WordDocumentBuilder:
    """Word文档构建器"""

    def __init__(self):
        self.document = Document()
        self.parser = MarkdownParser()
        self._setup_styles()
        self._setup_page_settings()

    def _setup_styles(self):
        styles = self.document.styles

        for i in range(1, 7):
            style_name = f"Heading {i}"
            if style_name in styles:
                style = styles[style_name]
                font = style.font
                font.bold = True
                font.size = Pt(20 - (i - 1) * 2)
                font.color.rgb = RGBColor(0, 0, 0)
                font.name = "等线"
                self._set_east_asian_font(style, "等线")

        normal_style = styles["Normal"]
        normal_style.font.name = "等线"
        normal_style.font.size = Pt(11)
        self._set_east_asian_font(normal_style, "等线")

    def _set_east_asian_font(self, style, font_name: str):
        try:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        except:
            pass

    def _set_run_font(self, run, font_name: str = "等线", font_size: int = None):
        run.font.name = font_name
        if font_size is not None:
            run.font.size = Pt(font_size)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _setup_page_settings(self):
        section = self.document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    def build(self, elements: List[MarkdownElement], images_dir: str = None):
        for element in elements:
            self._add_element(element, images_dir)

        self._add_footnotes_section()

    def _add_element(self, element: MarkdownElement, images_dir: str = None):
        handlers = {
            ElementType.HEADING: self._add_heading,
            ElementType.PARAGRAPH: self._add_paragraph,
            ElementType.CODE_BLOCK: self._add_code_block,
            ElementType.BLOCKQUOTE: self._add_blockquote,
            ElementType.LIST_ORDERED: self._add_ordered_list,
            ElementType.LIST_UNORDERED: self._add_unordered_list,
            ElementType.LIST_TASK: self._add_task_list,
            ElementType.TABLE: self._add_table,
            ElementType.HORIZONTAL_RULE: self._add_horizontal_rule,
            ElementType.MATH_BLOCK: self._add_math_block,
        }

        handler = handlers.get(element.element_type)
        if handler:
            handler(element, images_dir)

    def _add_heading(self, element: MarkdownElement, images_dir: str = None):
        heading_style = f"Heading {min(element.level, 6)}"
        paragraph = self.document.add_heading("", level=min(element.level, 6))
        self._add_runs_to_paragraph(
            paragraph, element.content, images_dir, is_heading=True
        )

    def _add_paragraph(self, element: MarkdownElement, images_dir: str = None):
        paragraph = self.document.add_paragraph()
        self._add_runs_to_paragraph(paragraph, element.content, images_dir)

    def _add_runs_to_paragraph(
        self,
        paragraph,
        runs: List[TextRun],
        images_dir: str = None,
        is_heading: bool = False,
    ):
        for run_data in runs:
            text = run_data.text

            if text.startswith("[IMAGE:"):
                parts = text[7:-1].split(":", 1)
                if len(parts) >= 2:
                    img_path = parts[0]
                    img_alt = parts[1] if len(parts) > 1 else ""
                    self._add_image(paragraph, img_path, img_alt, images_dir)
                continue

            run = paragraph.add_run(text)

            if is_heading:
                run.bold = True
            else:
                run.bold = run_data.bold
            run.italic = run_data.italic

            if run_data.strikethrough:
                run.font.strike = True

            if run_data.code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(192, 0, 0)
                self._set_run_shading(run, RGBColor(240, 240, 240))
                self._set_run_east_asian_font(run, "Consolas")
            else:
                if is_heading:
                    self._set_run_font(run, "等线", None)
                else:
                    self._set_run_font(run, "等线", 11)

            if run_data.link:
                self._add_hyperlink(paragraph, run, run_data.link)

    def _set_run_east_asian_font(self, run, font_name: str):
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _add_hyperlink(self, paragraph, run, url: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "等线")
        rFonts.set(qn("w:hAnsi"), "等线")
        rFonts.set(qn("w:eastAsia"), "等线")
        rPr.append(rFonts)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)

        new_run.append(rPr)

        text_elem = OxmlElement("w:t")
        text_elem.text = run.text
        new_run.append(text_elem)

        hyperlink.append(new_run)

        run._r.getparent().remove(run._r)
        paragraph._p.append(hyperlink)

    def _set_run_shading(self, run, color: RGBColor):
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), str(color))
        rPr.append(shd)

    def _add_image(
        self, paragraph, img_path: str, alt_text: str, images_dir: str = None
    ):
        if images_dir:
            full_path = os.path.join(images_dir, img_path)
        else:
            full_path = img_path

        if os.path.exists(full_path):
            try:
                run = paragraph.add_run()
                run.add_picture(full_path, width=Inches(5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                paragraph.add_run(f"[图片加载失败: {img_path}]")
        else:
            paragraph.add_run(f"[图片未找到: {img_path}]")

    def _add_code_block(self, element: MarkdownElement, images_dir: str = None):
        code_text = element.content
        language = element.language

        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        if language:
            lang_run = paragraph.add_run(f"语言: {language}\n")
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(128, 128, 128)
            lang_run.font.italic = True
            self._set_run_font(lang_run, "等线", 8)

        code_run = paragraph.add_run(code_text)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(9)
        self._set_run_east_asian_font(code_run, "Consolas")

        self._set_paragraph_shading(paragraph, RGBColor(245, 245, 245))

    def _set_paragraph_shading(self, paragraph, color: RGBColor):
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), str(color))
        pPr.append(shd)

    def _add_blockquote(self, element: MarkdownElement, images_dir: str = None):
        quote_lines = element.content

        for line_text in quote_lines:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(1)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            runs = self.parser._parse_inline(line_text)
            for run_data in runs:
                run = paragraph.add_run(run_data.text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(102, 102, 102)
                self._set_run_font(run, "等线", 11)

            self._add_left_border(paragraph, RGBColor(204, 204, 204))

    def _add_left_border(self, paragraph, color: RGBColor):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), str(color))
        pBdr.append(left)
        pPr.append(pBdr)

    def _add_ordered_list(self, element: MarkdownElement, images_dir: str = None):
        counter = 1
        last_indent = 0
        indent_counters = {}

        for item in element.content:
            indent = item.get("indent", 0)

            if indent not in indent_counters:
                indent_counters[indent] = 1
            else:
                indent_counters[indent] += 1

            for k in list(indent_counters.keys()):
                if k > indent:
                    del indent_counters[k]

            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5 + indent * 0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.5)

            number_run = paragraph.add_run(f"{indent_counters[indent]}. ")
            number_run.font.bold = True
            self._set_run_font(number_run, "等线", 11)

            self._add_runs_to_paragraph(paragraph, item["content"], images_dir)

    def _add_unordered_list(self, element: MarkdownElement, images_dir: str = None):
        bullets = ["•", "○", "■", "□", "◆", "◇"]

        for item in element.content:
            indent = item.get("indent", 0)
            bullet = bullets[indent % len(bullets)]

            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5 + indent * 0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.5)

            bullet_run = paragraph.add_run(f"{bullet} ")
            bullet_run.font.bold = True
            self._set_run_font(bullet_run, "等线", 11)

            self._add_runs_to_paragraph(paragraph, item["content"], images_dir)

    def _add_task_list(self, element: MarkdownElement, images_dir: str = None):
        for item in element.content:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.first_line_indent = Cm(-0.5)

            checked = item.get("checked", False)
            checkbox = "☑ " if checked else "☐ "
            checkbox_run = paragraph.add_run(checkbox)
            checkbox_run.font.size = Pt(12)
            self._set_run_font(checkbox_run, "等线", 12)

            self._add_runs_to_paragraph(paragraph, item["content"], images_dir)

    def _add_table(self, element: MarkdownElement, images_dir: str = None):
        table_data = element.content
        header = table_data["header"]
        rows = table_data["rows"]
        alignments = table_data["alignments"]

        num_cols = len(header)
        num_rows = len(rows) + 1

        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_row = table.rows[0]
        for idx, cell_text in enumerate(header):
            cell = header_row.cells[idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = self._get_alignment(
                    alignments[idx] if idx < len(alignments) else "left"
                )
                for run in paragraph.runs:
                    run.font.bold = True
                    self._set_run_font(run, "等线", 11)
            self._set_cell_shading(cell, RGBColor(230, 230, 230))

        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = self._get_alignment(
                        alignments[col_idx] if col_idx < len(alignments) else "left"
                    )
                    for run in paragraph.runs:
                        self._set_run_font(run, "等线", 11)

    def _get_alignment(self, align_str: str) -> WD_ALIGN_PARAGRAPH:
        alignments = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        return alignments.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

    def _set_cell_shading(self, cell, color: RGBColor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), str(color))
        tcPr.append(shd)

    def _add_horizontal_rule(self, element: MarkdownElement, images_dir: str = None):
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)

        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(pBdr)

    def _add_math_block(self, element: MarkdownElement, images_dir: str = None):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        math_run = paragraph.add_run(element.content)
        math_run.font.italic = True
        math_run.font.name = "Cambria Math"
        math_run.font.size = Pt(11)
        self._set_run_east_asian_font(math_run, "等线")

        self._set_paragraph_shading(paragraph, RGBColor(250, 250, 250))

    def _add_footnotes_section(self):
        pass

    def add_footnotes(self, footnotes: Dict[str, str]):
        if not footnotes:
            return

        self.document.add_paragraph()
        hr_paragraph = self.document.add_paragraph()
        hr_paragraph.paragraph_format.space_before = Pt(12)
        hr_paragraph.paragraph_format.space_after = Pt(6)

        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "000000")
        pBdr.append(top)

        pPr = hr_paragraph._p.get_or_add_pPr()
        pPr.append(pBdr)

        for footnote_id, footnote_text in footnotes.items():
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)

            id_run = paragraph.add_run(f"[^{footnote_id}]: ")
            id_run.font.bold = True
            id_run.font.size = Pt(10)
            self._set_run_font(id_run, "等线", 10)

            text_run = paragraph.add_run(footnote_text)
            text_run.font.size = Pt(10)
            self._set_run_font(text_run, "等线", 10)

    def save(self, output_path: str):
        self.document.save(output_path)


class MarkdownToWordConverter:
    """Markdown转Word转换器主类"""

    def __init__(self, images_dir: str = None):
        self.parser = MarkdownParser()
        self.builder = WordDocumentBuilder()
        self.images_dir = images_dir

    def convert(self, markdown_text: str, output_path: str):
        elements = self.parser.parse(markdown_text)
        self.builder.build(elements, self.images_dir)
        self.builder.add_footnotes(self.parser.footnotes)
        self.builder.save(output_path)

    def convert_file(self, input_path: str, output_path: str):
        with open(input_path, "r", encoding="utf-8") as f:
            markdown_text = f.read()

        input_dir = os.path.dirname(os.path.abspath(input_path))
        if self.images_dir is None:
            self.images_dir = input_dir

        self.convert(markdown_text, output_path)


def convert_markdown_to_word(input_path, output_path=None, images_dir=None):
    """
    把markdown文件转换为word文档(推荐使用),支持除了内嵌html以外的几乎所有语法,你可以先通过create_file工具创建markdown文件再调用此工具实现创建word文档

    参数:
        input_path (str): markdown文件路径
        output_path (str, optional): 输出word文件路径，如果为None则自动生成
        images_dir (str, optional): 图片目录路径，如果为None则为markdown文件的父目录
    返回:
        str: 执行状态信息
    """

    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}.docx"

    converter = MarkdownToWordConverter(images_dir=images_dir)

    try:
        converter.convert_file(input_path, output_path)
        return f"转换成功！输出文件: {output_path}"
    except Exception as e:
        return f"转换失败: {str(e)}"